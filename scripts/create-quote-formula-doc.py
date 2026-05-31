"""Generate RS Cleaning quote formula Word doc in Downloads."""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    raise SystemExit("python-docx required: pip install python-docx")

OUT = Path.home() / "Downloads" / "RS-Cleaning-Quote-Pricing-Formula.docx"

MIN_SQFT = 500
MIN_JOB = 125

doc = Document()
doc.add_heading("RS Cleaning Collective — Quote Pricing Formula", 0)
doc.add_paragraph(
    "This document describes how the website quote wizard and booking flow calculate estimated totals. "
    "Source of truth: assets/js/config.js, assets/js/quote-assistant.js, docs/QUOTE_PRICING.md."
)

doc.add_heading("Minimums", level=1)
doc.add_paragraph(f"Minimum square footage (typed entry only): {MIN_SQFT} sq ft.")
doc.add_paragraph(
    f"Minimum job (base price floor): ${MIN_JOB}. If basePrice is greater than $0 and less than ${MIN_JOB}, "
    f"basePrice is set to ${MIN_JOB}."
)
doc.add_paragraph(
    "Beds/baths estimates ignore the sq ft minimum (see formula below). Square checkout requires total ≥ $125."
)

doc.add_heading("Step 1 — Square footage", level=1)
doc.add_paragraph(f"If the customer enters square footage, that value is used (minimum {MIN_SQFT} sq ft).")
doc.add_paragraph("If using bedrooms / bathrooms:")
p = doc.add_paragraph()
p.add_run("sqft = round(beds × 450 + baths × 150 + 350)").bold = True
doc.add_paragraph("Defaults when missing: bedrooms → 2, bathrooms → 1.")
doc.add_paragraph("Example: 1 bed, 1 bath → 950 sq ft.")

doc.add_heading("Step 2 — Base price", level=1)
doc.add_paragraph("basePrice = sqft × rate[service type]")
table = doc.add_table(rows=5, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Service type"
hdr[1].text = "Rate ($ per sq ft)"
rows = [
    ("Standard cleaning", "0.17"),
    ("Deep cleaning", "0.28"),
    ("Move-in/Move-out", "0.32"),
    ("Office cleaning", "0.20"),
]
for i, (a, b) in enumerate(rows, 1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b
doc.add_paragraph(
    f"If basePrice > 0 and basePrice < ${MIN_JOB}, basePrice = ${MIN_JOB}."
)

doc.add_heading("Step 3 — Add-ons (flat fees, summed)", level=1)
t2 = doc.add_table(rows=9, cols=2)
t2.style = "Table Grid"
t2.rows[0].cells[0].text = "Add-on"
t2.rows[0].cells[1].text = "Price"
addons = [
    ("Wash and fold", "$45"),
    ("Fold laundry only", "$25"),
    ("Inside oven", "$40"),
    ("Inside fridge", "$40"),
    ("Cabinet interiors", "$50"),
    ("Interior Windows (1-10)", "$50"),
    ("Interior Windows (11-20)", "$100"),
    ("Bedding refresh", "$15"),
]
for i, (a, b) in enumerate(addons, 1):
    t2.rows[i].cells[0].text = a
    t2.rows[i].cells[1].text = b

doc.add_heading("Step 4 — Frequency discount", level=1)
doc.add_paragraph("subtotal = (basePrice + addonsPrice) × (1 − discount)")
t3 = doc.add_table(rows=5, cols=2)
t3.style = "Table Grid"
t3.rows[0].cells[0].text = "Frequency"
t3.rows[0].cells[1].text = "Discount"
for i, (a, b) in enumerate(
    [
        ("Weekly", "20%"),
        ("Bi-weekly", "15%"),
        ("Monthly", "10%"),
        ("One-time", "0%"),
    ],
    1,
):
    t3.rows[i].cells[0].text = a
    t3.rows[i].cells[1].text = b

doc.add_heading("Step 5 — Travel fee", level=1)
doc.add_paragraph(
    "Primary counties (Charles, St. Mary's, Calvert, Prince George's): $0. "
    "Outside / extended travel zone: $35 default. "
    "Resolved from Areas check, quote session, or booking address (flexible match)."
)

doc.add_heading("Step 6 — Final total", level=1)
p2 = doc.add_paragraph()
p2.add_run("total = round((subtotal + travelFee) × 100) / 100").bold = True

doc.add_heading("Example: 1 bed, 1 bath, standard, one-time, in-county", level=1)
doc.add_paragraph("Sq ft: 950 | Base: 950 × 0.17 = $161.50")
doc.add_paragraph("Add-ons: $0 | Discount: 0% | Travel: $0")
doc.add_paragraph("Estimated total: approximately $161.50")

doc.add_heading(f"Example: typed {MIN_SQFT} sq ft, standard, outside area", level=1)
doc.add_paragraph(f"Base: {MIN_SQFT} × 0.17 = $85 → floor ${MIN_JOB}")
doc.add_paragraph(f"After discount: ${MIN_JOB} | Travel: $35")
doc.add_paragraph(f"Estimated total: approximately ${MIN_JOB + 35}")

doc.save(OUT)
print(str(OUT))
